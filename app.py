from flask import Flask, request, send_file, render_template
from openpyxl import load_workbook
from docx import Document
import os

app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
TEMPLATE_FILE = "template.docx"

# Ensure upload folder exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


@app.route('/')
def upload_page():
    return render_template('upload.html')


@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return "No file uploaded", 400

    file = request.files['file']
    if file.filename == '':
        return "No selected file", 400

    # Save uploaded Excel file
    excel_path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(excel_path)

    # Process the Excel file and update the Word template
    updated_word_path = os.path.join(UPLOAD_FOLDER, "updated_document.docx")
    process_file(excel_path, TEMPLATE_FILE, updated_word_path)

    # Provide the updated Word document for download
    return send_file(updated_word_path, as_attachment=True)


def process_file(excel_file, word_template, output_file):
    from docx.shared import Pt  # For any required formatting adjustments

    # Load Excel data
    wb = load_workbook(excel_file)
    sheet = wb.active

    # Load Word template
    doc = Document(word_template)

    # Create a dictionary for placeholders and values
    data_mapping = {}
    for row in sheet.iter_rows(min_row=2, max_row=sheet.max_row, min_col=1, max_col=2, values_only=True):
        placeholder, value = row
        data_mapping[str(placeholder)] = str(value)

    # Function to replace text in a paragraph
    def replace_placeholder_in_paragraph(paragraph, mapping):
        for placeholder, value in mapping.items():
            if placeholder and paragraph.text and placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text:  # Only process paragraphs with text
            replace_placeholder_in_paragraph(paragraph, data_mapping)

    # Replace placeholders in tables (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholder_in_paragraph(paragraph, data_mapping)

    # Save the updated Word file
    doc.save(output_file)


if __name__ == '__main__':
    app.run(debug=True)
